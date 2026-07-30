import streamlit as st
import docx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import json
from google import genai
from PIL import Image
import datetime

st.set_page_config(page_title="Автоматизация ДКП и Актов", page_icon="🚗", layout="wide")

st.title("🚗 Генератор ДКП и Актов (AI-версия)")
st.caption("Умное распознавание документов через нейросеть Gemini с поддержкой кредитных сделок")

# --- БОКОВАЯ ПАНЕЛЬ: ЗАГРУЗКА И AI ---
st.sidebar.header("⚙️ Настройки и загрузка")

api_key = st.sidebar.text_input("🔑 Вставьте API-ключ Gemini", type="password", help="Начинается на AIzaSy...")

uploaded_files = st.sidebar.file_uploader(
    "Фото (Паспорт, СТС, ПТС)",
    type=["jpg", "jpeg", "png"],
    accept_multiple_files=True
)

# --- ИНИЦИАЛИЗАЦИЯ ДАННЫХ (ЖЕЛЕЗОБЕТОННОЕ СОХРАНЕНИЕ) ---
today_str = datetime.datetime.now().strftime("%d.%m.%Y")

DEFAULT_DATA = {
    # Стандартные данные
    "contract_num": "АК00000589",
    "contract_date": f"{today_str} г.",
    "city": "г. Калининград",
    "buyer_company": "ООО «Авто-К»",
    "buyer_rep": "Руководителя отдела продаж Маркова Андрея Дмитриевича",
    "buyer_poa": "Доверенности б/н от 10.06.2026",
    "buyer_address": "236003, Калининградская обл., г. Калининград, Московский пр-кт, д. 250, помещение 1",
    "buyer_phone": "8 (4012) 777-213",
    "buyer_inn": "3906203848",
    "buyer_kpp": "390601001",
    "buyer_ogrn": "1083925041310",
    "buyer_bank": "КБ \"ЭНЕРГОТРАНСБАНК\" (АО)",
    "buyer_rs": "40702810800000011431",
    "buyer_ks": "30101810800000000701",
    "buyer_bik": "042748701",
    "seller_fio": "",
    "seller_passport": "",
    "seller_address": "",
    "seller_phone": "",
    "car_mark": "",
    "car_vin": "",
    "car_engine": "Отсутствует",
    "car_body": "Отсутствует",
    "car_frame": "Отсутствует",
    "car_color": "",
    "car_mileage": "",
    "car_pts": "",
    "car_sts": "",
    "car_number": "",
    "car_year": "",
    "car_defects": "диагностика прилагается",
    "price_num": "1 500 000",
    "price_text": "Один миллион пятьсот тысяч рублей",
    "payment_details": "денежные средства перечислить в качестве частичной оплаты приобретаемого Продавцом нового автомобиля",
    
    # Кредит и залог (Переключатель)
    "is_credit_deal": False,
    "pledge_bank_full": "ПАО «Сбербанк» Россия, Москва, 117312, ул. Вавилова, д. 19",
    "pledge_contract": "14100686811",
    "pledge_date": "27.03.2024г.",
    "pledge_bank_short": "ПАО «Сбербанк» Россия, Москва",
    "pledge_debt_num": "886810,79",
    "pledge_debt_text": "Восемьсот восемьдесят шесть тысяч восемьсот десять рублей 79 копейки",
    "pledge_repay_date": f"{today_str} г.",
    "tradein_amount_num": "613000",
    "tradein_amount_text": "Шестьсот тринадцать тысяч рублей",
    "new_car_name": "BelGee x50+",
    "new_car_vin": "Y4K8622Z0TB927651",
    "new_car_contract": "2026-305",
    "new_car_date": f"{today_str} г.",
    "credit_pay_num": "887000",
    "credit_pay_text": "восемьсот восемьдесят семь тысяч рублей",
    "credit_pay_date": "31.07.2026г",
    "seller_acc_num": "40817810120861423605",
    "seller_bank_name": "КАЛИНИНГРАДСКОЕ ОТДЕЛЕНИЕ N8626 ПАО СБЕРБАНК",
    "seller_bik": "042748634",
    "seller_ks": "30101810100000000634",
    "seller_bank_inn": "7707083893",
    "seller_bank_kpp": "390643005",
    "seller_bank_okpo": "09134363",
    "seller_bank_ogrn": "1027700132195",
    "seller_swift": "SABRRU2P920",
    "seller_bank_address1": "236006, КАЛИНИНГРАД, МОСКОВСКИЙ ПР., 24",
    "seller_bank_address2": "236000, Г. КАЛИНИНГРАД, УЛ. КАРЛА МАРКСА, Д.49, 51 / УЛ. КОММУНАЛЬНАЯ, Д. 29-35"
}

for k, v in DEFAULT_DATA.items():
    if k not in st.session_state:
        st.session_state[k] = v

# --- КНОПКА СКАНИРОВАНИЯ ---
if st.sidebar.button("🤖 Умное распознавание", type="primary"):
    if not api_key.startswith("AIzaSy"):
        st.sidebar.error("⚠️ Введите правильный API-ключ (начинается на AIzaSy...)")
    elif not uploaded_files:
        st.sidebar.warning("⚠️ Загрузите фото документов.")
    else:
        with st.spinner("Нейросеть читает документы... (занимает 5-15 сек)"):
            try:
                client = genai.Client(api_key=api_key.strip())
                images = [Image.open(f) for f in uploaded_files]
                
                prompt = """
                Твоя задача — извлечь данные из предоставленных фото (паспорт, СТС, ПТС) и вернуть их СТРОГО в формате JSON без разметки markdown.
                Структура JSON:
                {
                  "seller_fio": "Фамилия Имя Отчество (если есть в паспорте)",
                  "seller_passport": "Серия и номер, кем выдан, дата выдачи, код подразделения (собери все в одну строку)",
                  "seller_address": "Адрес регистрации (прописка) полностью",
                  "car_mark": "Марка и модель авто",
                  "car_vin": "VIN номер авто (17 символов)",
                  "car_year": "Год выпуска",
                  "car_pts": "Серия и номер ПТС",
                  "car_sts": "Серия и номер СТС",
                  "car_number": "Государственный регистрационный знак авто",
                  "car_color": "Цвет авто",
                  "car_engine": "Модель и номер двигателя (если есть)"
                }
                Если какого-то поля на фото нет, оставь пустую строку "".
                """
                
                response = client.models.generate_content(
                    model='gemini-2.0-flash',
                    contents=[prompt, *images]
                )
                
                cleaned_text = response.text.replace("```json", "").replace("```", "").strip()
                extracted_data = json.loads(cleaned_text)
                
                for key, val in extracted_data.items():
                    if key in st.session_state and val:
                        st.session_state[key] = str(val)
                
                if st.session_state.get("car_vin"):
                     st.session_state["car_body"] = st.session_state["car_vin"]
                        
                st.sidebar.success("✅ Все данные успешно извлечены!")
                st.rerun()
            except Exception as e:
                st.sidebar.error(f"❌ Ошибка: {e}")

# --- ИНТЕРФЕЙС ВВОДА ДАННЫХ ---
st.subheader("📋 Данные для ДКП и Акта")
tab_deal, tab_credit, tab_seller, tab_buyer, tab_car = st.tabs(["📄 Договор и Сумма", "🏦 Залог / Кредит", "👤 Продавец", "🏢 Покупатель", "🚘 Автомобиль"])

with tab_deal:
    c1, c2, c3 = st.columns(3)
    c1.text_input("Номер Договора", key="contract_num")
    c2.text_input("Дата", key="contract_date")
    c3.text_input("Город", key="city")
    
    c4, c5 = st.columns([1, 2])
    c4.text_input("Общая сумма (цифрами)", key="price_num")
    c5.text_input("Общая сумма (прописью)", key="price_text")
    st.text_area("Особые условия оплаты (п. 2.2 для стандарта)", key="payment_details", height=80, help="Используется только если выключена галочка кредита на соседней вкладке.")

with tab_credit:
    st.checkbox("🔥 Автомобиль в залоге (Сложный расчет: трейд-ин + погашение)", key="is_credit_deal")
    
    if st.session_state.is_credit_deal:
        st.markdown("##### 📌 Данные банка и залога (Пункт 1.2)")
        cr1, cr2 = st.columns(2)
        cr1.text_area("Банк (полный адрес)", key="pledge_bank_full", height=70)
        cr1.text_input("Банк (короткое название)", key="pledge_bank_short")
        cr2.text_input("№ Кредитного договора", key="pledge_contract")
        cr2.text_input("Дата кредитного договора", key="pledge_date")
        
        cr3, cr4, cr5 = st.columns(3)
        cr3.text_input("Сумма долга (цифрами)", key="pledge_debt_num")
        cr4.text_input("Сумма долга (прописью)", key="pledge_debt_text")
        cr5.text_input("Дата погашения", key="pledge_repay_date")

        st.markdown("##### 📌 Сложный расчет оплаты (Пункт 2.2)")
        st.caption("Часть 1: Оплата в счет нового авто")
        tr1, tr2 = st.columns(2)
        tr1.text_input("Сумма Трейд-ин (цифрами)", key="tradein_amount_num")
        tr1.text_input("Сумма Трейд-ин (прописью)", key="tradein_amount_text")
        tr2.text_input("Новое авто (Марка)", key="new_car_name")
        tr2.text_input("Новое авто (VIN)", key="new_car_vin")
        
        tr3, tr4 = st.columns(2)
        tr3.text_input("№ Договора на новое авто", key="new_car_contract")
        tr4.text_input("Дата договора на новое авто", key="new_car_date")

        st.caption("Часть 2: Перечисление остатка в банк")
        bp1, bp2, bp3 = st.columns(3)
        bp1.text_input("Остаток кредита (цифрами)", key="credit_pay_num")
        bp2.text_input("Остаток кредита (пропись)", key="credit_pay_text")
        bp3.text_input("Перечислить не позднее", key="credit_pay_date")
        
        br1, br2 = st.columns(2)
        br1.text_input("Номер счёта", key="seller_acc_num")
        br1.text_input("Банк получателя", key="seller_bank_name")
        br1.text_input("БИК", key="seller_bik")
        br1.text_input("Корр. счёт", key="seller_ks")
        br1.text_input("SWIFT-код", key="seller_swift")
        
        br2.text_input("ИНН банка", key="seller_bank_inn")
        br2.text_input("КПП банка", key="seller_bank_kpp")
        br2.text_input("ОКПО", key="seller_bank_okpo")
        br2.text_input("ОГРН", key="seller_bank_ogrn")
        
        st.text_input("Почтовый адрес банка", key="seller_bank_address1")
        st.text_input("Почтовый адрес доп.офиса", key="seller_bank_address2")

with tab_seller:
    st.text_input("ФИО Продавца", key="seller_fio")
    st.text_area("Паспорт (серия, номер, кем/когда выдан)", key="seller_passport", height=100)
    st.text_input("Адрес регистрации", key="seller_address")
    st.text_input("Телефон Продавца", key="seller_phone")

with tab_buyer:
    c1, c2 = st.columns(2)
    c1.text_input("Название организации (Покупатель)", key="buyer_company")
    c1.text_area("В лице (Должность, ФИО)", key="buyer_rep")
    c1.text_input("Действующего на основании", key="buyer_poa")
    
    c2.text_area("Юр. Адрес", key="buyer_address")
    c2.text_input("ИНН", key="buyer_inn")
    c2.text_input("КПП", key="buyer_kpp")
    c2.text_input("ОГРН", key="buyer_ogrn")
    
    c3, c4 = st.columns(2)
    c3.text_input("Банк", key="buyer_bank")
    c3.text_input("Р/С", key="buyer_rs")
    c4.text_input("К/С", key="buyer_ks")
    c4.text_input("БИК", key="buyer_bik")

with tab_car:
    c1, c2, c3 = st.columns(3)
    c1.text_input("Марка (Модель)", key="car_mark")
    c2.text_input("VIN", key="car_vin")
    c3.text_input("Год выпуска", key="car_year")
    
    c4, c5, c6 = st.columns(3)
    c4.text_input("Цвет", key="car_color")
    c5.text_input("Гос. номер", key="car_number")
    c6.text_input("Пробег", key="car_mileage")
    
    c7, c8 = st.columns(2)
    c7.text_input("ПТС", key="car_pts")
    c8.text_input("СТС", key="car_sts")
    
    c9, c10, c11 = st.columns(3)
    c9.text_input("№ Двигателя", key="car_engine")
    c10.text_input("№ Кузова", key="car_body")
    c11.text_input("№ Рамы (Шасси)", key="car_frame")
    st.text_input("Дефекты (недостатки)", key="car_defects")


# --- ФУНКЦИЯ ГЕНЕРАЦИИ WORD ---
def create_word_doc(data):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    
    # Заголовок ДКП
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"ДОГОВОР КУПЛИ-ПРОДАЖИ АВТОМОБИЛЯ № {data['contract_num']}").bold = True
    
    p = doc.add_paragraph()
    p.add_run(f"{data['city']}").bold = False
    p.add_run(f"\t\t\t\t\t\t\t\t\t\t\t\t{data['contract_date']}")
    
    doc.add_paragraph(
        f"{data['buyer_company']}, в лице {data['buyer_rep']}, действующего(ей) на основании {data['buyer_poa']}, "
        f"именуемый(ая) в дальнейшем «Покупатель», с одной стороны, и {data['seller_fio']}, именуемый(ая) "
        f"в дальнейшем «Продавец», с другой стороны, совместно именуемые «Стороны» заключили настоящий договор о нижеследующем:"
    )
    
    p = doc.add_paragraph()
    p.add_run("1. Предмет Договора").bold = True
    doc.add_paragraph(
        "1.1. Продавец обязуется передать в собственность Покупателя, а Покупатель обязуется принять и оплатить следующий "
        "бывший в эксплуатации автомобиль (далее по тексту - «Автомобиль»):\n"
        f"Марка (Модель): {data['car_mark']};\n"
        f"VIN: {data['car_vin']};\n"
        f"Модель, № двигателя: {data['car_engine']};\n"
        f"Номер кузова: {data['car_body']};\n"
        f"Номер рамы (шасси): {data['car_frame']};\n"
        f"Цвет: {data['car_color']};\n"
        f"Пробег (по показаниям одометра): {data['car_mileage']};\n"
        f"ПТС: {data['car_pts']};\n"
        f"Свидетельство о регистрации: {data['car_sts']};\n"
        f"Государственный регистрационный номер: {data['car_number']};\n"
        f"Год выпуска: {data['car_year']};\n"
        f"Дефекты (недостатки): {data['car_defects']}."
    )
    
    # --- ЛОГИКА 1.2 ЗАЛОГ vs СТАНДАРТ ---
    if data['is_credit_deal']:
        doc.add_paragraph(f"1.2. Продаваемый по настоящему договору автомобиль находится в залоге у {data['pledge_bank_full']} по кредитному договору №{data['pledge_contract']} от {data['pledge_date']}. Сумма задолженности Продавца перед {data['pledge_bank_short']} {data['pledge_debt_num']} рублей ({data['pledge_debt_text']}) на момент полного досрочного погашения кредита {data['pledge_repay_date']}.")
    else:
        doc.add_paragraph("1.2. Продавец гарантирует, что Автомобиль принадлежит ему на праве индивидуальной собственности, не заложен, не арендован, не является предметом спора, не состоит под арестом, не числится в розыске и не обременен никакими обязательствами перед третьими лицами.")
    
    doc.add_paragraph("1.3. Продавец гарантирует, что при заключении настоящего Договора уведомил Покупателя обо всех известных ему дефектах Автомобиля.")
    doc.add_paragraph("1.4. Продавец гарантирует, что на Автомобиле оригинальный пробег, и все техническое обслуживание осуществлялось надлежащим образом.")
    doc.add_paragraph("1.5. В случае нарушения Продавцом гарантий, Покупатель вправе в одностороннем порядке расторгнуть настоящий Договор.")
    
    p = doc.add_paragraph()
    p.add_run("2. Цена товара и порядок оплаты").bold = True
    doc.add_paragraph(f"2.1. Общая сумма Договора (стоимость Автомобиля) составляет {data['price_num']} рублей ({data['price_text']}), НДС не предусмотрен.")
    
    # --- ЛОГИКА 2.2 СЛОЖНЫЙ РАСЧЕТ vs СТАНДАРТ ---
    if data['is_credit_deal']:
        doc.add_paragraph(f"2.2. Расчет по договору осуществляется следующим образом:\n- денежные средства в размере {data['tradein_amount_num']} ({data['tradein_amount_text']}), за автомобиль перечислить в качестве частичной оплаты приобретаемого Продавцом ({data['seller_fio']}) у {data['buyer_company']} нового автомобиля {data['new_car_name']} VIN {data['new_car_vin']} по договору №{data['new_car_contract']} от {data['new_car_date']} по заявлению;\n- частичная стоимость автомобиля в размере {data['credit_pay_num']} рублей ({data['credit_pay_text']}) перечисляются не позднее {data['credit_pay_date']} на счет:\nНомер счёта: {data['seller_acc_num']}\nБанк получателя: {data['seller_bank_name']}\nБИК: {data['seller_bik']}\nКорр. счёт: {data['seller_ks']}\nИНН: {data['seller_bank_inn']}\nКПП: {data['seller_bank_kpp']}\nОКПО: {data['seller_bank_okpo']}\nОГРН: {data['seller_bank_ogrn']}\nSWIFT-код: {data['seller_swift']}\nПочтовый адрес банка: {data['seller_bank_address1']}\nПочтовый адрес доп.офиса: {data['seller_bank_address2']}.")
    else:
        doc.add_paragraph(f"2.2. Расчет по договору осуществляется следующим образом:\n- {data['payment_details']}.")
    
    p = doc.add_paragraph()
    p.add_run("3. Передача Автомобиля").bold = True
    doc.add_paragraph("3.1. Автомобиль передается Покупателю в день подписания настоящего Договора.")
    doc.add_paragraph("3.2. Передача Автомобиля Покупателю оформляется актом приема-передачи, подписываемым Сторонами.")
    doc.add_paragraph("3.3. При приемке Автомобиля Покупатель за счет Продавца осуществляет его проверку.")
    doc.add_paragraph("3.4. Продавец передает Покупателю комплект документов, относящихся к Автомобилю.")
    doc.add_paragraph("3.5. Право собственности на Автомобиль переходит от Продавца к Покупателю с момента подписания Акта приема-передачи.")
    doc.add_paragraph("3.6. Продавец уведомлен о том, что Покупатель не регистрирует на свое имя транспортные средства, предназначенные для продажи.")
    
    p = doc.add_paragraph()
    p.add_run("4. Прочие условия").bold = True
    doc.add_paragraph("4.1. Обстоятельства непреодолимой силы освобождают Стороны от обязательств на время действия таких обстоятельств.")
    doc.add_paragraph("4.2. В случае, если выяснится, что Автомобиль является предметом спора, Продавец обязуется компенсировать убытки Покупателя.")
    doc.add_paragraph("4.3. Все споры будут по возможности решаться путем переговоров.")
    doc.add_paragraph("4.4. В случае невозможности разрешения разногласий, спор подлежит рассмотрению в суде.")
    doc.add_paragraph("4.5. Настоящий Договор вступает в силу с момента его подписания.")
    doc.add_paragraph("4.6. Настоящий Договор составлен в трех экземплярах, два экземпляра — для Покупателя, один — для Продавца.")

    p = doc.add_paragraph()
    p.add_run("5. Использование и передача персональных данных").bold = True
    doc.add_paragraph("5.1. Подписывая настоящий Договор Продавец, сведения о котором содержатся в настоящем Договоре и иных документах переданных и/или представленных Покупателю, дает свое согласие на обработку Покупателем, включая получение от Продавца и/или от любых третьих лиц, с учетом требований действующего законодательства РФ, персональных данных Продавца и подтверждает, что дает такое согласие, действуя своей волей и в своем интересе. Согласие дается Продавцом для целей заключения с Продавцом любых договоров, их дальнейшего исполнения, участия в проводимых Покупателем акциях, опросах, исследованиях (включая, но, не ограничиваясь, проведением опросов посредством электронной, телефонной и сотовой связи), или совершения иных действий, порождающих юридические последствия в отношении Продавца или других лиц, предоставления Продавцу информации об оказываемых Покупателем услугах. Перечень персональных данных, на обработку которых дается согласие: фамилия, имя, отчество; год, месяц, дата рождения; пол; адрес; контактный телефон (домашний, рабочий, мобильный); контактный адрес электронной почты; сведения о профессиональной деятельности; марка, модель, VIN, государственный номерной знак приобретенного/обслуживаемого автомобиля; пробег приобретенного/обслуживаемого автомобиля; название дилерского центра, где приобретен/обслуживался/ ремонтировался автомобиль; дата выдачи автомобиля при покупке/из сервиса; перечень работ, проведенных с автомобилем; перечень замененных на автомобиле деталей и прочая информация.")
    doc.add_paragraph("5.2. Обработка персональных данных Продавца осуществляется Покупателем в объеме, который необходим для достижения каждой из вышеперечисленных целей. Продавец подтверждает, что данное согласие действует в течение срока хранения Покупателем персональных данных Продавца, составляющего 10 (десять) лет с момента их получения, который продлятся на каждые следующие 10 (десять) лет при условии отсутствия у Покупателя сведений об отзыве настоящего согласия. Продавец вправе отозвать свое согласие путем направления соответствующего письменного уведомления Покупателю не менее чем за 3 (три) месяца до момента отзыва согласия. После получения такого уведомления Покупатель в течение трех рабочих дней прекращает обработку персональных данных. В случае отзыва согласия Продавца на обработку персональных данных, Покупатель вправе не прекращать обработку персональных данных и не уничтожать их в случаях, предусмотренных действующим законодательством РФ.")
    doc.add_paragraph("5.3. Настоящее согласие предоставляется для осуществления любых действий в отношении персональных данных Продавца, которые необходимы или желаемы для достижения каждой из указанных выше целей, включая без ограничений: сбор, систематизацию, накопление, хранение, уточнение. Продавец признаёт и подтверждает, что в случае необходимости предоставления персональных данных для достижения указанных выше целей третьему лицу, а равно как при привлечении третьих лиц к оказанию услуг, при передаче Покупателем принадлежащих ему функций и полномочий иному лицу, Покупатель вправе в необходимом объеме раскрывать для совершения вышеуказанных действий информацию с соблюдением требований законодательства. Также Продавец признает и подтверждает, что согласие считается данным им любым третьим лицам, указанным выше с учетом соответствующих изменений, и любые третьи лица имеют право на обработку персональных данных на основании настоящего согласия.")
    doc.add_paragraph(f"5.4. Настоящим Продавец выражает согласие и разрешает Покупателю передавать свои персональные данные, указанные в настоящем Договоре, следующим лицам при условии соблюдения ими требований законодательства РФ об обеспечении ими конфиденциальности и безопасности персональных данных при их обработке: {data['buyer_company']} (место нахождения: Адрес: {data['buyer_address']}).")
    doc.add_paragraph("5.5. Покупатель и третьи лица, которым переданы персональные данные, вправе обрабатывать персональные данные Продавца посредством внесения их в электронную базу данных, включения в списки (реестры) и отчетные формы, предусмотренные документами, регламентирующими предоставление отчетных данных (документов), и передавать их уполномоченным органам.")
    doc.add_paragraph("5.6. Продавец настоящим выражает согласие и разрешает Покупателю и третьим лицам обрабатывать свои персональные данные, в том числе информировать Продавца и получать информацию от Покупателя и третьих лиц с помощью различных средств связи, включая, но, не ограничиваясь: почтовая рассылка, электронная почта, телефон (включая мобильный), sms-рассылка, факсимильная связь, сеть Интернет. Продавец выражает свое согласие и разрешает Покупателю и третьим лицам осуществлять все действия, необходимые для достижения целей обработки персональных данных.")
    doc.add_paragraph("5.7. Настоящим Продавец уведомлен о том, что предполагаемыми пользователями персональных данных являются работники операторов (Покупателя, Общество с ограниченной ответственностью «Акамацу Моторс» и третьих лиц), а также лица, привлеченные на условиях гражданско-правового договора. Правовым основанием обработки персональных данных третьими лицами является наличие договорных отношений.")

    p = doc.add_paragraph()
    p.add_run("6. Юридические адреса и банковские реквизиты Сторон").bold = True
    
    table = doc.add_table(rows=1, cols=2)
    row = table.rows[0]
    row.cells[0].text = f"ПРОДАВЕЦ:\n{data['seller_fio']}\nПаспорт: {data['seller_passport']}\nАдрес: {data['seller_address']}\nТел. {data['seller_phone']}\n\n________________________/____________/"
    row.cells[1].text = f"ПОКУПАТЕЛЬ:\n{data['buyer_company']}\nАдрес: {data['buyer_address']}\nТелефон: {data['buyer_phone']}\nИНН {data['buyer_inn']} КПП {data['buyer_kpp']}\nОГРН {data['buyer_ogrn']}\nБанк: {data['buyer_bank']}\nР/С {data['buyer_rs']}\nК/С {data['buyer_ks']}\nБИК {data['buyer_bik']}\n\n________________________/____________/"
    
    doc.add_paragraph("\nУважаемый Клиент! Благодарим Вас за выбор нашей компании. Обращаем Ваше внимание, что Вы вправе прекратить регистрацию Автомобиля в органах ГИБДД по своему заявлению по истечению 10 дней с момента перехода права собственности на Автомобиль.").italic = True
    
    doc.add_page_break()
    
    # --- АКТ ПРИЕМА ПЕРЕДАЧИ ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("АКТ ПРИЕМА–ПЕРЕДАЧИ").bold = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"К ДОГОВОРУ КУПЛИ-ПРОДАЖИ № {data['contract_num']} от {data['contract_date']}")
    
    p = doc.add_paragraph()
    p.add_run(f"{data['city']}").bold = False
    p.add_run(f"\t\t\t\t\t\t\t\t\t\t\t\t{data['contract_date']}")
    
    doc.add_paragraph(
        f"{data['buyer_company']} в лице {data['buyer_rep']} действующего(ей) на основании {data['buyer_poa']}, "
        f"именуемый(ая) в дальнейшем «Покупатель», с одной стороны, и {data['seller_fio']}, именуемый(ая) "
        f"в дальнейшем «Продавец», с другой стороны, совместно именуемые «Стороны» заключили настоящий Акт о нижеследующем:"
    )
    
    doc.add_paragraph(
        f"1. Действуя в рамках Договора купли-продажи автомобиля № {data['contract_num']} от {data['contract_date']}, "
        f"заключенного между сторонами, Продавец передал, а Покупатель принял следующий товар:\n\n"
        f"Марка (Модель): {data['car_mark']};\n"
        f"VIN: {data['car_vin']};\n"
        f"Модель, № двигателя: {data['car_engine']};\n"
        f"Номер кузова: {data['car_body']};\n"
        f"Номер рамы (шасси): {data['car_frame']};\n"
        f"Цвет: {data['car_color']};\n"
        f"Пробег: {data['car_mileage']};\n"
        f"ПТС: {data['car_pts']};\n"
        f"Свидетельство о регистрации: {data['car_sts']};\n"
        f"Государственный регистрационный номер: {data['car_number']};\n"
        f"Год выпуска: {data['car_year']};\n"
        f"Дефекты (недостатки): {data['car_defects']}.\n"
    )
    
    doc.add_paragraph(f"2. Стоимость товара составляет {data['price_num']} рублей ({data['price_text']}), НДС не предусмотрен.")
    doc.add_paragraph("3. Вместе с товаром Продавец передал Покупателю:\n- Паспорт транспортного средства (ПТС)\n- Свидетельство регистрации ТC\n- Ключи от а/м\n- Дополнительный комплект резины\n- Комплект ковриков в салон")
    
    table_act = doc.add_table(rows=1, cols=2)
    row_act = table_act.rows[0]
    row_act.cells[0].text = f"Продавец\n\n_______________________/ {data['seller_fio']} /"
    row_act.cells[1].text = f"Покупатель\n\n______________________/ Представитель /"
    
    target = io.BytesIO()
    doc.save(target)
    target.seek(0)
    return target

# --- КНОПКА ГЕНЕРАЦИИ И СКАЧИВАНИЯ ---
st.divider()

st.download_button(
    label="📄 СГЕНЕРИРОВАТЬ И СКАЧАТЬ ПОЛНЫЙ КОМПЛЕКТ (.docx)",
    data=create_word_doc(st.session_state),
    file_name=f"ДКП_Акт_{st.session_state['car_mark']}_{st.session_state['car_vin']}.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    type="primary",
    use_container_width=True
)
